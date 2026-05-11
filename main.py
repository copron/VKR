from pathlib import Path
from typing import Dict, Any

import joblib
import pandas as pd
import streamlit as st
from docx import Document
import time


# Конфиг путей
BASE_DIR = Path(__file__).resolve().parent
TFIDF_PATH = BASE_DIR / "tfidf_vectorizer.pkl"
MODELS_PATH = BASE_DIR / "criterion_models_best.pkl"


CRITERIA_MAX = {
    "criterion_01": 2, "criterion_02": 1, "criterion_03": 1, "criterion_04": 2,
    "criterion_05": 1, "criterion_06": 1, "criterion_07": 3, "criterion_08": 1,
    "criterion_09": 2, "criterion_10": 1, "criterion_11": 3, "criterion_12": 3,
    "criterion_13": 2, "criterion_14": 1, "criterion_15": 1, "criterion_16": 2,
    "criterion_17": 2, "criterion_18": 1, "criterion_19": 2, "criterion_20": 2,
    "criterion_21": 1, "criterion_22": 2,
}


CRITERION_TO_STAGE = {
    "criterion_01": "pre_analytics",
    "criterion_02": "pre_analytics",
    "criterion_03": "pre_analytics",
    "criterion_04": "pre_analytics",
    "criterion_05": "test_design",
    "criterion_06": "test_design",
    "criterion_07": "test_design",
    "criterion_08": "test_design",
    "criterion_09": "monitoring",
    "criterion_10": "monitoring",
    "criterion_11": "tech_launch",
    "criterion_12": "tech_launch",
    "criterion_13": "results_summary",
    "criterion_14": "results_summary",
    "criterion_15": "results_summary",
    "criterion_16": "results_summary",
    "criterion_17": "results_summary",
    "criterion_18": "results_summary",
    "criterion_19": "presentation",
    "criterion_20": "presentation",
    "criterion_21": "presentation",
    "criterion_22": "presentation",
}


STAGE_ORDER = [
    "pre_analytics",
    "test_design",
    "monitoring",
    "tech_launch",
    "results_summary",
    "presentation",
]


STAGE_LABELS = {
    "pre_analytics": "1. Преданалитика",
    "test_design": "2. Дизайн теста",
    "monitoring": "3. Мониторинг",
    "tech_launch": "4. Тех. запуск",
    "results_summary": "5. Итоги",
    "presentation": "6. Презентация",
}


# Подгрузка моделей
@st.cache_resource(show_spinner=True)
def load_tfidf_and_models():
    tfidf = joblib.load(TFIDF_PATH)

    if not MODELS_PATH.exists():
        raise FileNotFoundError(f"Не найден файл моделей: {MODELS_PATH}")

    models = joblib.load(MODELS_PATH)

    if not isinstance(models, dict):
        raise ValueError("criterion_models_best.pkl должен содержать dict {criterion: model}")

    return tfidf, models


# Чтение DOCX
def read_docx_text_simple(file) -> str:
    if isinstance(file, (str, Path)):
        doc = Document(str(file))
    else:
        doc = Document(file)

    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + table_lines)
    return full_text


# Валидация текста документа
def normalize_text_for_check(text: str) -> str:
    return " ".join(str(text).lower().replace("\n", " ").split())


def validate_document_text(raw_text: str, min_len: int = 300):
    text = normalize_text_for_check(raw_text)

    if not text:
        return False, "Не удалось извлечь текст из файла: документ пустой или повреждён."

    if len(text) < min_len:
        return False, (
            f"Документ содержит слишком мало текста для автоматической оценки "
            f"(меньше {min_len} символов)."
        )

    return True, None


def is_likely_ab_test_card(raw_text: str):
    text = normalize_text_for_check(raw_text)

    required_sections = [
        "цель эксперимента",
        "общая информация",
    ]

    missing_sections = [section for section in required_sections if section not in text]
    is_card = len(missing_sections) == 0

    debug_info = {
        "required_sections": required_sections,
        "missing_sections": missing_sections,
    }

    return is_card, debug_info


# Предсказание и агрегирование
def predict_scores_from_text(raw_text: str, tfidf_obj, models_dict: Dict[str, Any]) -> Dict[str, int]:
    X = tfidf_obj.transform([str(raw_text)])
    per_crit: Dict[str, int] = {}

    for crit in CRITERIA_MAX.keys():
        clf = models_dict.get(crit)

        if crit == "criterion_05":
            # этот критерий по смыслу всегда считаем выполненным, если карточка есть
            per_crit[crit] = 1
        elif clf is not None:
            y_pred = clf.predict(X)[0]
            per_crit[crit] = int(y_pred)
        else:
            # на всякий случай дефолт для критериев без модели
            per_crit[crit] = 0

    return per_crit


def build_stage_and_total(per_crit: Dict[str, int]):
    rows = []
    for crit, score in per_crit.items():
        maxscore = CRITERIA_MAX.get(crit, None)
        stage = CRITERION_TO_STAGE.get(crit, "unknown")
        rows.append({
            "stage": stage,
            "criterion": crit,
            "score": score,
            "maxscore": maxscore,
        })

    df = pd.DataFrame(rows)

    df["maxscore"] = df["maxscore"].fillna(0).astype(int)

    total_score = int(df["score"].sum())
    total_max = int(df["maxscore"].sum())
    total_pct = (total_score / total_max * 100) if total_max > 0 else 0.0

    stage_df = (
        df.groupby("stage", as_index=False)
        .agg(
            stage_score=("score", "sum"),
            stage_max=("maxscore", "sum"),
        )
    )

    stage_df["stage_pct"] = stage_df.apply(
        lambda r: (r["stage_score"] / r["stage_max"] * 100) if r["stage_max"] > 0 else 0.0,
        axis=1,
    )

    stage_df["stage_label"] = stage_df["stage"].map(STAGE_LABELS).fillna(stage_df["stage"])

    stage_df["stage_order"] = stage_df["stage"].apply(
        lambda s: STAGE_ORDER.index(s) if s in STAGE_ORDER else 999
    )
    stage_df = stage_df.sort_values("stage_order").reset_index(drop=True)

    crit_df = df.copy()
    crit_df["stage_label"] = crit_df["stage"].map(STAGE_LABELS).fillna(crit_df["stage"])

    return stage_df, crit_df, total_score, total_max, total_pct


# Подсветка DataFrame
def color_row_by_score(row, score_col: str, max_col: str):
    score = row[score_col]
    maxscore = row[max_col]

    if maxscore <= 0:
        return [""] * len(row)

    if score <= 0:
        return ["background-color: rgba(255, 0, 0, 0.10)"] * len(row)
    elif score < maxscore:
        return ["background-color: rgba(255, 215, 0, 0.10)"] * len(row)
    else:
        return ["background-color: rgba(0, 128, 0, 0.10)"] * len(row)


def summarize_result(total_score: int, total_max: int, stage_df: pd.DataFrame) -> str:
    if total_max <= 0:
        return "Не удалось посчитать итоговую оценку — отсутствуют максимальные баллы."

    pct = total_score / total_max * 100

    if pct >= 85:
        level = "высокое качество проведения A/B"
    elif pct >= 60:
        level = "среднее качество проведения A/B"
    else:
        level = "низкое качество проведения A/B"

    weak_stages = stage_df[stage_df["stage_pct"] < 100]["stage_label"].tolist()

    if weak_stages:
        weak_text = ", ".join(weak_stages)
        advice = f"Рекомендуется доработать этапы: {weak_text}."
    else:
        advice = "Все этапы оформлены на максимальный балл."

    return f"Общий результат: {level} (≈{pct:.1f}% от максимума). {advice}"


# UI
def main():
    st.set_page_config(
        page_title="Система оценки качества проведения A/B теста",
        layout="wide"
    )

    st.title("Система оценки качества проведения A/B теста")

    st.markdown("""
    <style>
    div[data-testid="stMetric"] {
        background-color: #f8f9fa;
        border: 1px solid #e9ecef;
        padding: 16px 20px;
        border-radius: 12px;
    }
    div[data-testid="stMetricValue"] {
        font-size: 40px;
        font-weight: 700;
    }
    div[data-testid="stMetricLabel"] {
        font-size: 18px;
        font-weight: 600;
    }
    </style>
    """, unsafe_allow_html=True)

    with st.spinner("Загружаем TF-IDF и модели..."):
        tfidf, models_linearsvc = load_tfidf_and_models()
    st.success("Модели загружены успешно! Ждём Ваш документ с карточкой...")

    st.markdown("---")

    st.subheader("1. Загрузите DOCX с карточкой теста")
    uploaded_file = st.file_uploader("Выберите файл .docx", type=["docx"])

    if uploaded_file is None:
        st.info("Загрузите файл, чтобы увидеть оценку.")
        return

    if not uploaded_file.name.lower().endswith(".docx"):
        st.error("Ошибка: система принимает только файлы формата .docx.")
        return

    st.write(f"Файл: **{uploaded_file.name}**")

    with st.spinner("Читаем DOCX и извлекаем текст..."):
        try:
            card_text = read_docx_text_simple(uploaded_file)
        except Exception:
            st.error("Ошибка: не удалось прочитать файл как документ Word (.docx).")
            return

    is_valid_text, validation_error = validate_document_text(card_text, min_len=300)
    if not is_valid_text:
        st.error(validation_error)
        return

    is_card_like, card_debug = is_likely_ab_test_card(card_text)
    if not is_card_like:
        st.error(
            "Документ не похож на карточку A/B-теста: "
            "не найдены обязательные разделы 'Цель эксперимента' и/или 'Общая информация'."
        )
        with st.expander("Проверка обязательных разделов", expanded=False):
            st.write("Обязательные разделы:")
            st.write(", ".join(card_debug["required_sections"]))
            st.write("Не найдены:")
            st.write(", ".join(card_debug["missing_sections"]))
        return

    with st.expander("Показать извлечённый текст", expanded=False):
        st.text(card_text[:4000])

    st.markdown("---")

    st.subheader("2. Оценка по этапам и критериям")

    if st.button("Посчитать оценку", type="primary"):
        with st.spinner("Считаем ML-оценки по критериям..."):
            t0 = time.perf_counter()
            per_crit = predict_scores_from_text(card_text, tfidf, models_linearsvc)
            stage_df, crit_df, total_score, total_max, total_pct = build_stage_and_total(per_crit)
            pred_time_sec = time.perf_counter() - t0

        col_metric_1, col_metric_2 = st.columns(2)

        with col_metric_1:
            st.metric(
                label="Итоговая оценка",
                value=f"{total_score} / {total_max}",
                delta=f"{total_pct:.1f}% от максимума"
            )

        with col_metric_2:
            st.metric(
                label="Время расчёта",
                value=f"{pred_time_sec:.4f} c"
            )

        st.markdown("#### Оценка по этапам")
        stage_view = stage_df[["stage_label", "stage_score", "stage_max", "stage_pct"]].copy()
        stage_view = stage_view.rename(columns={
            "stage_label": "Этап",
            "stage_score": "Оценка этапа системой",
            "stage_max": "Максимальный балл за этап",
            "stage_pct": "% этапа",
        })

        stage_styled = stage_view.style.apply(
            color_row_by_score,
            axis=1,
            score_col="Оценка этапа системой",
            max_col="Максимальный балл за этап",
        )
        st.dataframe(stage_styled, use_container_width=True)

        st.markdown("#### Оценка по критериям")
        crit_view = crit_df[["stage_label", "criterion", "score", "maxscore"]].copy()
        crit_view = crit_view.rename(columns={
            "stage_label": "Этап",
            "criterion": "Критерий",
            "score": "Оценка критерия системой",
            "maxscore": "Максимальный балл за критерий",
        })
        crit_view = crit_view.sort_values(["Этап", "Критерий"]).reset_index(drop=True)

        crit_styled = crit_view.style.apply(
            color_row_by_score,
            axis=1,
            score_col="Оценка критерия системой",
            max_col="Максимальный балл за критерий",
        )
        st.dataframe(crit_styled, use_container_width=True, height=500)

        st.markdown("#### Cаммари")
        summary_text = summarize_result(total_score, total_max, stage_df)
        st.write(summary_text)


if __name__ == "__main__":
    main()